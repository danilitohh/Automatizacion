"""Parser de documentos que conserva orden, tablas y relaciones básicas."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from pypdf import PdfReader
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook

from .models import SemanticDocument, SemanticNode
from .normalizer import display_text, normalized_text


class DocumentParser:
    """Convierte DOCX, XLSX y texto plano en nodos semánticos genéricos."""

    def parse(self, filename: str, content: bytes) -> SemanticDocument:
        extension = Path(filename or "").suffix.lower()
        if extension == ".docx":
            return self._parse_docx(content, filename)
        if extension == ".xlsx":
            return self._parse_xlsx(content, filename)
        if extension == ".pdf":
            return self._parse_pdf(content, filename)
        if extension in {".txt", ".md", ".csv"}:
            return self._parse_text(content, filename)
        raise ValueError("Formato no soportado. Usa .pdf, .docx, .xlsx, .txt, .md o .csv.")

    def _parse_docx(self, content: bytes, filename: str) -> SemanticDocument:
        try:
            document = Document(BytesIO(content))
        except Exception as error:
            raise ValueError("No se pudo leer el documento fuente DOCX.") from error

        nodes: list[SemanticNode] = []
        section = ""
        has_title = False
        order = 0
        for block in self._iter_blocks(document):
            order += 1
            if isinstance(block, Paragraph):
                text = display_text(block.text)
                if not text:
                    continue
                node_type = self._paragraph_type(block, text)
                group_match = re.match(r"^(\d+\s*[°º]\s*(?:cuatrimestre|semestre|trimestre|periodo|bloque|módulo))(?:\s+(.+))?$", text, re.IGNORECASE)
                if group_match:
                    section = display_text(group_match.group(1))
                    nodes.append(SemanticNode(f"doc-{order}-group", "subsection", section, section, order, {"block": order}, {"ordered_group": True}))
                    if not group_match.group(2):
                        continue
                    text = display_text(group_match.group(2))
                    node_type = "paragraph"
                elif node_type == "paragraph" and not has_title:
                    node_type = "title"
                    has_title = True
                if node_type in {"title", "subtitle", "section", "subsection"}:
                    section = text
                metadata: dict[str, Any] = {"style": block.style.name if block.style else ""}
                label, value = self._label_value(text)
                if label:
                    node_type = "label_value"
                    metadata.update({"label": label, "value": value})
                if text.endswith("?") or text.startswith(("¿", "Q:", "Pregunta:")):
                    node_type = "question"
                nodes.append(SemanticNode(f"doc-{order}", node_type, text, section, order, {"block": order}, metadata))
            else:
                table_node = self._table_node(block, order, section)
                if table_node:
                    nodes.append(table_node)

        title = next((node.text for node in nodes if node.type == "title"), "")
        return SemanticDocument("document", title, nodes, {"filename": filename, "format": "docx"})

    def _parse_xlsx(self, content: bytes, filename: str) -> SemanticDocument:
        try:
            workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        except Exception as error:
            raise ValueError("No se pudo leer el documento fuente XLSX.") from error
        nodes: list[SemanticNode] = []
        order = 0
        for sheet in workbook.worksheets:
            for row_number, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                values = [display_text(str(value or "")) for value in row]
                expanded: list[str] = []
                for value in values:
                    expanded.extend(part for part in (display_text(item) for item in value.splitlines()) if part)
                values = expanded
                values = [value for value in values if value]
                if not values:
                    continue
                order += 1
                if len(values) == 2:
                    nodes.append(SemanticNode(
                        f"doc-{order}", "label_value", " | ".join(values), sheet.title, order,
                        {"sheet": sheet.title, "row": row_number}, {"label": values[0], "value": values[1]},
                    ))
                else:
                    nodes.append(SemanticNode(f"doc-{order}", "table_row", " | ".join(values), sheet.title, order, {"sheet": sheet.title, "row": row_number}, {"cells": values}))
        workbook.close()
        return SemanticDocument("document", "", nodes, {"filename": filename, "format": "xlsx"})

    def _parse_text(self, content: bytes, filename: str) -> SemanticDocument:
        text = content.decode("utf-8-sig", errors="replace")
        nodes: list[SemanticNode] = []
        section = ""
        for order, raw_line in enumerate(text.splitlines(), start=1):
            line = display_text(raw_line)
            if not line:
                continue
            node_type = "list_item" if re.match(r"^(?:[-*•]|\d+[.)])\s+", line) else "paragraph"
            if line.startswith("#"):
                node_type = "section" if not line.startswith("##") else "subsection"
                line = line.lstrip("# ")
                section = line
            label, value = self._label_value(line)
            metadata = {"label": label, "value": value} if label else {}
            if label:
                node_type = "label_value"
            nodes.append(SemanticNode(f"doc-{order}", node_type, line, section, order, {"line": order}, metadata))
        title = next((node.text for node in nodes if node.type in {"title", "section"}), "")
        return SemanticDocument("document", title, nodes, {"filename": filename, "format": Path(filename).suffix.lower()})

    def _parse_pdf(self, content: bytes, filename: str) -> SemanticDocument:
        try:
            reader = PdfReader(BytesIO(content))
        except Exception as error:
            raise ValueError("No se pudo leer el documento fuente PDF.") from error
        nodes: list[SemanticNode] = []
        order = 0
        for page_number, page in enumerate(reader.pages, start=1):
            for raw_line in (page.extract_text() or "").splitlines():
                line = display_text(raw_line)
                if not line:
                    continue
                order += 1
                node_type = "list_item" if re.match(r"^(?:[-*â€¢]|\d+[.)])\s+", line) else "paragraph"
                nodes.append(SemanticNode(f"pdf-{order}", node_type, line, f"PÃ¡gina {page_number}", order, {"page": page_number}))
        title = nodes[0].text if nodes else ""
        return SemanticDocument("document", title, nodes, {"filename": filename, "format": ".pdf", "pages": len(reader.pages)})

    @staticmethod
    def _iter_blocks(document: DocumentType) -> Iterable[Paragraph | Table]:
        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, document)

    def _table_node(self, table: Table, order: int, section: str) -> SemanticNode | None:
        rows: list[list[str]] = []
        for row in table.rows:
            values: list[str] = []
            for cell in row.cells:
                cell_values = [display_text(value) for value in cell.text.splitlines() if display_text(value)]
                values.extend(cell_values or [display_text(cell.text)])
            values = [value for value in values if value]
            if values:
                rows.append(values)
        if not rows:
            return None
        children: list[SemanticNode] = []
        for index, values in enumerate(rows, start=1):
            metadata: dict[str, Any] = {"cells": values}
            if len(values) >= 2:
                metadata.update({"label": values[0], "value": " | ".join(values[1:])})
            children.append(SemanticNode(f"doc-{order}-{index}", "table_row", " | ".join(values), section, index, {"table": order, "row": index}, metadata))
        return SemanticNode(f"doc-{order}", "table", "", section, order, {"table": order}, {"rows": len(rows)}, children)

    @staticmethod
    def _paragraph_type(paragraph: Paragraph, text: str) -> str:
        style = (paragraph.style.name if paragraph.style else "").lower()
        if "title" in style:
            return "title"
        if "heading 1" in style:
            return "section"
        if "heading" in style:
            return "subsection"
        if re.match(r"^\d+\.\s+[^?]{3,100}$", text) and not text.endswith((".", ":")):
            return "section"
        if "list" in style or re.match(r"^(?:[-*•]|\d+[.)])\s+", text):
            return "list_item"
        if len(text) <= 100 and text.isupper():
            return "section"
        return "paragraph"

    @staticmethod
    def _label_value(text: str) -> tuple[str, str]:
        match = re.match(r"^([^:|]{2,80})\s*:\s*(.+)$", text)
        return (display_text(match.group(1)), display_text(match.group(2))) if match else ("", "")
