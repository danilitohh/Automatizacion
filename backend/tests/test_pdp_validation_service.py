"""Pruebas de la lectura de fuentes del módulo PDP vs DOCX."""

from io import BytesIO

from docx import Document
from openpyxl import Workbook

from backend.app.config.settings import Settings
from backend.app.services.pdp_validation_service import PdpValidationService


def build_excel() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "PDP"
    sheet.append(["Carrera", "URL PDP"])
    sheet.append(["Licenciatura en Ciencia de Datos", "https://ejemplo.test/ciencia-de-datos"])
    content = BytesIO()
    workbook.save(content)
    return content.getvalue()


def build_docx() -> bytes:
    document = Document()
    document.add_heading("Licenciatura en Ciencia de Datos", level=1)
    document.add_heading("Descripción", level=2)
    document.add_paragraph("Forma especialistas para analizar datos y resolver problemas de negocio.")
    document.add_heading("Asignaturas", level=2)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Periodo"
    table.cell(0, 1).text = "Asignatura"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "Estadística Aplicada"
    document.add_heading("Preguntas frecuentes", level=2)
    document.add_paragraph("¿Cuál es la duración del programa?")
    content = BytesIO()
    document.save(content)
    return content.getvalue()


def test_reads_and_maps_excel_and_docx_sources(tmp_path):
    settings = Settings(database_path=tmp_path / "pdp.db", storage_dir=tmp_path / "storage")
    service = PdpValidationService(settings)

    programs, excel_info = service._read_excel(build_excel())
    blocks = service._read_docx(build_docx())
    references = service._map_document_to_programs(programs, blocks)
    reference = references[service._normalize("Licenciatura en Ciencia de Datos")]

    assert excel_info["programs_loaded"] == 1
    assert programs[0].url == "https://ejemplo.test/ciencia-de-datos"
    assert reference.found is True
    assert reference.description == ["Forma especialistas para analizar datos y resolver problemas de negocio."]
    assert reference.subjects == ["Estadística Aplicada"]
    assert reference.faqs == ["¿Cuál es la duración del programa?"]


def test_compares_sections_with_page_text(tmp_path):
    service = PdpValidationService(Settings(database_path=tmp_path / "pdp.db", storage_dir=tmp_path / "storage"))
    page_text = """
        Licenciatura en Ciencia de Datos. Forma especialistas para analizar datos y resolver problemas de negocio.
        Asignaturas: Estadística Aplicada. Preguntas frecuentes: ¿Cuál es la duración del programa?
    """

    subjects = service._compare_items(["Estadística Aplicada"], page_text, "Asignaturas")
    description = service._compare_section(
        ["Forma especialistas para analizar datos y resolver problemas de negocio."], page_text, "Descripción"
    )

    assert subjects["status"] == "PASS"
    assert description["status"] == "PASS"
