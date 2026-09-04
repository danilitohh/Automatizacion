"""Pruebas del pipeline genÃ©rico, sin nombres de programas concretos."""

from io import BytesIO

from docx import Document

from backend.app.modules.pdp_validation.comparison_engine import ComparisonEngine
from backend.app.modules.pdp_validation.document_parser import DocumentParser
from backend.app.modules.pdp_validation.models import SemanticDocument, SemanticNode
from backend.app.modules.pdp_validation.normalizer import normalized_text


def test_document_parser_detects_structure_and_tables():
    document = Document()
    document.add_heading("Producto demo", level=1)
    document.add_heading("Beneficios", level=2)
    document.add_paragraph("Nombre del producto")
    document.add_paragraph("Una descripciÃ³n general.")
    document.add_paragraph("Beneficio A", style="List Bullet")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "DuraciÃ³n"
    table.rows[0].cells[1].text = "24 meses"
    extra_row = table.add_row()
    extra_row.cells[0].text = "1"
    extra_row.cells[1].text = "Materia A\nMateria B"

    parsed = DocumentParser().parse("fuente.docx", BytesIO(_save(document)).getvalue())

    types = {node.type for node in parsed.nodes}
    assert {"title", "subsection", "paragraph", "list_item", "table"}.issubset(types)
    assert parsed.nodes[-1].children[0].metadata["label"] == "DuraciÃ³n"
    assert any("Materia A" in child.metadata["cells"] for child in parsed.nodes[-1].children)


def test_comparison_reports_normalized_missing_extra_and_duplicate():
    expected = SemanticDocument("document", nodes=[
        SemanticNode("e1", "section", "Beneficios", "Beneficios"),
        SemanticNode("e2", "list_item", "Educaci" + chr(243) + "n ambiental", "Beneficios"),
        SemanticNode("e3", "list_item", "Acompa" + chr(241) + "amiento", "Beneficios"),
    ])
    actual = SemanticDocument("web", nodes=[
        SemanticNode("a1", "section", "Beneficios", "Beneficios"),
        SemanticNode("a2", "list_item", chr(8226) + " educacion ambiental", "Beneficios"),
        SemanticNode("a3", "list_item", "Acompa" + chr(241) + "amiento", "Beneficios"),
        SemanticNode("a4", "list_item", "Acompa" + chr(241) + "amiento", "Beneficios"),
        SemanticNode("a5", "paragraph", "Precio especial", "Beneficios"),
    ])

    comparison = ComparisonEngine().compare(expected, actual)
    statuses = {finding["status"] for finding in comparison["findings"]}

    assert "MATCH_NORMALIZADO" in statuses
    assert "DUPLICADO" in statuses
    assert "EXTRA" in statuses


def test_normalizer_ignores_question_numbering_and_table_prefixes():
    assert normalized_text("2. " + "La carrera tiene validez oficial?") == normalized_text("La carrera tiene validez oficial?")
    expected = SemanticDocument("document", nodes=[SemanticNode("row", "table_row", "", "Plan", metadata={"cells": ["1", "Fundamentos de la educaci" + chr(243) + "n"]})])
    actual = SemanticDocument("web", nodes=[SemanticNode("item", "list_item", "Fundamentos de la educaci" + chr(243) + "n", "Asignaturas")])
    findings = ComparisonEngine().compare(expected, actual)["findings"]
    assert any(finding["status"] == "MATCH_EXACTO" for finding in findings)


def _save(document):
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
